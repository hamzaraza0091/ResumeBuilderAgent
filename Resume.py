import streamlit as st
from groq import Groq
import os
from dotenv import load_dotenv
import io
import re

load_dotenv()

# Page configuration
st.set_page_config(
    page_title="Resume Builder Assistant",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .stApp {
        background-color: #0a0a0a;
    }
    
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #ffffff;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    
    .sub-header {
        font-size: 1rem;
        color: #a1a1aa;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    .chat-message {
        padding: 1rem;
        border-radius: 0.75rem;
        margin-bottom: 1rem;
    }
    
    .user-message {
        background-color: #18181b;
        border: 1px solid #27272a;
    }
    
    .assistant-message {
        background-color: #09090b;
        border: 1px solid #27272a;
    }
    
    .sidebar-section {
        background-color: #18181b;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    
    .resume-preview {
        background-color: #ffffff;
        color: #000000;
        padding: 2rem;
        border-radius: 0.5rem;
        font-family: 'Georgia', serif;
    }
    
    .stTextInput > div > div > input {
        background-color: #18181b;
        border: 1px solid #27272a;
        color: #ffffff;
    }
    
    .stButton > button {
        background-color: #2563eb;
        color: white;
        border: none;
        border-radius: 0.5rem;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    
    .stButton > button:hover {
        background-color: #1d4ed8;
    }
    
    div[data-testid="stSidebarContent"] {
        background-color: #09090b;
    }

    /* Download section styling */
    .download-section {
        background-color: #18181b;
        border: 1px solid #27272a;
        border-radius: 0.75rem;
        padding: 1.25rem;
        margin-top: 1.5rem;
    }

    .download-title {
        color: #ffffff;
        font-weight: 600;
        font-size: 1rem;
        margin-bottom: 0.75rem;
    }

    div[data-testid="stDownloadButton"] > button {
        width: 100%;
        background-color: #16a34a;
        color: white;
        border: none;
        border-radius: 0.5rem;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }

    div[data-testid="stDownloadButton"] > button:hover {
        background-color: #15803d;
    }
</style>
""", unsafe_allow_html=True)

# System prompt for the Resume Builder Assistant
SYSTEM_PROMPT = """You are an intelligent Resume Builder Assistant.
Your job is to help users create professional resumes based on their provided information.
You must:
1. Ask the user for missing details such as:
   - Full Name
   - Contact Information
   - Career Objective
   - Education
   - Skills
   - Work Experience
   - Projects
   - Certifications
   - Achievements
2. Use conversation memory to remember previously provided information so the user doesn't need to repeat details.
3. If the user gives incomplete information, ask follow-up questions.
4. Format the final resume professionally with proper sections:
   - Personal Information
   - Summary
   - Education
   - Skills
   - Experience
   - Projects
   - Certifications
   - Achievements
5. Suggest improvements to weak resume descriptions.
6. Tailor resumes for specific job roles if the user asks 
   (e.g., Software Engineer, Data Analyst, AI Engineer).
7. Keep responses concise and professional.

If the user says "Generate my resume" or similar, create the final formatted resume using all stored conversation data.

When generating the final resume, format it cleanly with clear section headers and bullet points.
Use professional language and action verbs for experience descriptions.
"""

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []

if "resume_data" not in st.session_state:
    st.session_state.resume_data = {
        "name": "",
        "contact": "",
        "objective": "",
        "education": [],
        "skills": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "achievements": []
    }

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("GROQ_API_KEY", "")

if "last_resume_text" not in st.session_state:
    st.session_state.last_resume_text = ""


# ── Helper: detect if a message looks like a generated resume ──────────────
def _looks_like_resume(text: str) -> bool:
    """Return True if the assistant message appears to be a full resume."""
    resume_keywords = [
        "personal information", "contact information", "career objective",
        "professional summary", "work experience", "education", "skills",
        "projects", "certifications", "achievements"
    ]
    lower = text.lower()
    matches = sum(1 for kw in resume_keywords if kw in lower)
    return matches >= 3


# ── PDF generator (reportlab) ──────────────────────────────────────────────
def generate_pdf(resume_text: str) -> bytes:
    """Convert plain-text resume into a styled PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        accent = colors.HexColor("#2563eb")

        name_style = ParagraphStyle(
            "ResumeName",
            parent=styles["Title"],
            fontSize=20,
            textColor=colors.HexColor("#1e1e2e"),
            spaceAfter=4,
        )
        section_style = ParagraphStyle(
            "SectionHeader",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=accent,
            spaceBefore=10,
            spaceAfter=2,
        )
        body_style = ParagraphStyle(
            "ResumeBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=2,
        )
        bullet_style = ParagraphStyle(
            "ResumeBullet",
            parent=body_style,
            leftIndent=14,
            bulletIndent=0,
        )

        story = []
        lines = resume_text.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            # Detect markdown-style headers  ## or #
            if stripped.startswith("## ") or stripped.startswith("# "):
                header_text = stripped.lstrip("#").strip()
                story.append(Paragraph(header_text, section_style))
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=accent, spaceAfter=4))
            # ALL-CAPS lines as section headers
            elif stripped.isupper() and len(stripped) > 3 and not stripped.startswith("-"):
                story.append(Paragraph(stripped, section_style))
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=accent, spaceAfter=4))
            # Bullet points
            elif stripped.startswith(("- ", "• ", "* ")):
                bullet_text = stripped.lstrip("-•* ").strip()
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
            # Bold lines (potential name or sub-header)
            elif stripped.startswith("**") and stripped.endswith("**"):
                inner = stripped.strip("*")
                story.append(Paragraph(f"<b>{inner}</b>", body_style))
            else:
                # Strip leftover markdown bold
                clean = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", stripped)
                story.append(Paragraph(clean, body_style))

        doc.build(story)
        return buf.getvalue()

    except ImportError:
        return b""


# ── DOCX generator (python-docx) ───────────────────────────────────────────
def generate_docx(resume_text: str) -> bytes:
    """Convert plain-text resume into a styled DOCX file."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        accent_rgb = RGBColor(0x25, 0x63, 0xEB)

        def add_section_header(text):
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = accent_rgb
            # Bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2563EB")
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)

        def add_body(text):
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # strip bold markers
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(clean)
            run.font.size = Pt(10)
            return p

        def add_bullet(text):
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", text.lstrip("-•* ").strip())
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(clean)
            run.font.size = Pt(10)

        lines = resume_text.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue

            if stripped.startswith("## ") or stripped.startswith("# "):
                add_section_header(stripped.lstrip("#").strip())
            elif stripped.isupper() and len(stripped) > 3 and not stripped.startswith("-"):
                add_section_header(stripped)
            elif stripped.startswith(("- ", "• ", "* ")):
                add_bullet(stripped)
            else:
                add_body(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        return b""


def get_groq_response(messages: list, api_key: str) -> str:
    """Get response from Groq API."""
    try:
        client = Groq(api_key=api_key)
        full_messages = [{"role": "system", "content": SYSTEM_PROMPT}] + messages
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=full_messages,
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"


def display_chat_message(role: str, content: str):
    if role == "user":
        with st.chat_message("user", avatar="👤"):
            st.markdown(content)
    else:
        with st.chat_message("assistant", avatar="📄"):
            st.markdown(content)


# ── Sidebar ────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### Settings")

    api_key_input = st.text_input(
        "Groq API Key",
        type="password",
        value=st.session_state.api_key,
        placeholder="gsk_...",
        help="Enter your Groq API key to use the assistant"
    )
    if api_key_input:
        st.session_state.api_key = api_key_input

    st.divider()

    st.markdown("### Quick Actions")
    col1, col2 = st.columns(2)

    with col1:
        if st.button("Generate Resume", use_container_width=True):
            if st.session_state.api_key:
                st.session_state.messages.append({
                    "role": "user",
                    "content": "Generate my resume"
                })
                st.rerun()

    with col2:
        if st.button("Clear Chat", use_container_width=True):
            st.session_state.messages = []
            st.session_state.last_resume_text = ""
            st.session_state.resume_data = {
                "name": "", "contact": "", "objective": "",
                "education": [], "skills": [], "experience": [],
                "projects": [], "certifications": [], "achievements": []
            }
            st.rerun()

    # ── Download section (appears once a resume has been generated) ────────
    if st.session_state.last_resume_text:
        st.divider()
        st.markdown("### ⬇️ Download Resume")

        resume_txt = st.session_state.last_resume_text

        # Plain-text download (always available)
        st.download_button(
            label="📄 Download as TXT",
            data=resume_txt.encode("utf-8"),
            file_name="resume.txt",
            mime="text/plain",
            use_container_width=True,
        )

        # PDF download
        try:
            from reportlab.lib.pagesizes import letter  # noqa: F401
            pdf_bytes = generate_pdf(resume_txt)
            if pdf_bytes:
                st.download_button(
                    label="🔴 Download as PDF",
                    data=pdf_bytes,
                    file_name="resume.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
        except ImportError:
            st.caption("Install `reportlab` for PDF downloads.")

        # DOCX download
        try:
            from docx import Document  # noqa: F401
            docx_bytes = generate_docx(resume_txt)
            if docx_bytes:
                st.download_button(
                    label="🔵 Download as DOCX",
                    data=docx_bytes,
                    file_name="resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        except ImportError:
            st.caption("Install `python-docx` for Word downloads.")

    st.divider()

    st.markdown("### Tips")
    st.markdown("""
- Start by providing your **name** and **contact info**
- Share your **education** background
- List your **skills** and **experience**
- Mention any **projects** or **certifications**
- Say **"Generate my resume"** when ready
- Ask to tailor for specific roles
""")

    st.divider()

    st.markdown("### Example Prompts")
    example_prompts = [
        "My name is John Doe and I'm a software engineer",
        "I have 3 years of experience with Python and JavaScript",
        "I graduated from MIT with a CS degree in 2020",
        "Can you tailor my resume for a Data Scientist role?",
        "Generate my resume"
    ]
    for prompt in example_prompts:
        if st.button(prompt, key=f"example_{prompt[:20]}", use_container_width=True):
            if st.session_state.api_key:
                st.session_state.messages.append({"role": "user", "content": prompt})
                st.rerun()


# ── Main content ───────────────────────────────────────────────────────────
st.markdown('<h1 class="main-header">Resume Builder Assistant</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Create professional resumes through conversation</p>', unsafe_allow_html=True)

if not st.session_state.api_key:
    st.warning("Please enter your Groq API key in the sidebar to get started.")
    st.info("You can get an API key from [Groq Console](https://console.groq.com/keys)")
    st.stop()

# Display existing chat messages
for message in st.session_state.messages:
    display_chat_message(message["role"], message["content"])

# Generate assistant response when last message is from user
if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    with st.chat_message("assistant", avatar="📄"):
        with st.spinner("Thinking..."):
            response = get_groq_response(
                st.session_state.messages,
                st.session_state.api_key
            )
            st.markdown(response)

    st.session_state.messages.append({"role": "assistant", "content": response})

    # If the response looks like a resume, cache it for downloading
    if _looks_like_resume(response):
        st.session_state.last_resume_text = response
        st.rerun()  # refresh sidebar to show download buttons

# Chat input
if prompt := st.chat_input("Tell me about yourself or ask for help with your resume..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    st.rerun()

# Welcome message
if not st.session_state.messages:
    st.markdown("""
    <div style="text-align: center; padding: 2rem; background-color: #18181b; border-radius: 0.75rem; border: 1px solid #27272a;">
        <h3 style="color: #ffffff; margin-bottom: 1rem;">Welcome to Resume Builder Assistant!</h3>
        <p style="color: #a1a1aa;">I'll help you create a professional resume. Start by telling me:</p>
        <ul style="color: #a1a1aa; text-align: left; display: inline-block; margin-top: 1rem;">
            <li>Your name and contact information</li>
            <li>Your education background</li>
            <li>Your work experience</li>
            <li>Your skills and projects</li>
        </ul>
        <p style="color: #a1a1aa; margin-top: 1rem;">When you're ready, just say <strong>"Generate my resume"</strong>!</p>
        <p style="color: #71717a; margin-top: 0.5rem; font-size: 0.875rem;">📥 Download buttons (TXT · PDF · DOCX) appear in the sidebar once your resume is generated.</p>
    </div>
    """, unsafe_allow_html=True)